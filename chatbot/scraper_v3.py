import re
import json
import hashlib
from pathlib import Path
from urllib.parse import urljoin, urlparse, urldefrag
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
import pandas as pd
import trafilatura
from bs4 import BeautifulSoup
from tqdm import tqdm
from pypdf import PdfReader
from docx import Document
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

SITEMAP_URL = "https://mpr.gov.ba/bs/mapa-stranica"
OUTPUT_DIR = Path("mpr_dataset_v3_fixed")
FILES_DIR = OUTPUT_DIR / "files"
OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
FILES_DIR.mkdir(exist_ok=True, parents=True)

MAX_WORKERS = 8
TIMEOUT = 25
MIN_TEXT_LENGTH = 120
CHUNK_SIZE = 1100
CHUNK_OVERLAP = 180

DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".rtf"}

HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; MPR Sitemap Scraper Fixed; student-project)"
}

# Iz sitemap-a: uzimamo samo glavne sadržajne linkove do Projekti/Strategije.
# U web prikazu linkovi 4-180 odgovaraju Ministarstvo + Oblasti rada + Zakoni/Ugovori.
MIN_LINK_ID = 4
MAX_LINK_ID = 180

DROP_URL_PARTS = [
    "/hr", "/en", "/sr",
    "portal/objava",
    "categorySlug=vijesti",
    "tender",
    "javne-nabavke",
    "javni-oglasi",
    "konkurs",
    "galerija",
    "foto",
    "video"
]

DROP_TITLES = {
    "Hrvatski", "Bosanski", "Српски", "English", "Početna",
    "Projekti/Strategije", "Publikacije/Priručnici",
    "Tenderi/Javni oglasi", "Info/Pristup informacijama"
}

CATEGORY_RULES = [
    ("obrasci", ["obrazac", "obrasci", "formular", "formulari", "zahtjev", "prijava"]),
    ("registracija", ["registracija", "registracije", "udruženja", "udruzenja", "fondacije", "nvo", "registar udruženja", "registar udruzenja", "upis promjena", "brisanje iz registra"]),
    ("ispiti", ["pravosudni ispit", "stručni upravni ispit", "strucni upravni ispit", "ispit", "ispiti", "polaganje", "termini"]),
    ("pravna_pomoc", ["pravna pomoć", "pravna pomoc", "alimentacija", "otmica djece", "međunarodna pravna", "medjunarodna pravna"]),
    ("zakoni_i_propisi", ["zakon", "zakoni", "pravilnik", "pravilnici", "propisi", "podzakonski", "ugovori", "konvencije", "ustav"]),
    ("registri", ["registar", "registri", "evidencija", "izvod iz registra"]),
    ("kontakt_i_nadleznosti", ["kontakt", "kontakti", "nadležnosti", "nadleznosti", "organizacija", "sektor", "ured", "ministar"]),
]

def make_session():
    s = requests.Session()
    s.headers.update(HEADERS)

    retries = Retry(
        total=3,
        backoff_factor=0.6,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"]
    )

    adapter = HTTPAdapter(max_retries=retries, pool_connections=50, pool_maxsize=50)
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    return s

session = make_session()

def clean_text(text):
    if not text:
        return ""
    text = str(text).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def norm(text):
    text = clean_text(text).lower()
    return (
        text.replace("č", "c")
        .replace("ć", "c")
        .replace("š", "s")
        .replace("ž", "z")
        .replace("đ", "dj")
    )

def normalize_url(url):
    url, _ = urldefrag(url)
    p = urlparse(url)

    if p.scheme not in ["http", "https"]:
        return None

    if not p.netloc.endswith("mpr.gov.ba"):
        return None

    return url.rstrip("/")

def should_drop(url, title):
    u = norm(url)
    t = clean_text(title)

    if t in DROP_TITLES:
        return True

    return any(part in u for part in DROP_URL_PARTS)

def get_ext(url):
    return Path(urlparse(url).path).suffix.lower()

def is_document_url(url):
    return get_ext(url) in DOCUMENT_EXTENSIONS

def hash_text(text):
    return hashlib.md5(clean_text(text).encode("utf-8")).hexdigest()

def safe_filename(url):
    p = urlparse(url)
    name = Path(p.path).name or hashlib.md5(url.encode()).hexdigest()
    name = re.sub(r"[^a-zA-Z0-9_.-]", "_", name)

    if "." not in name:
        name += ".html"

    qhash = hashlib.md5((p.query or "").encode()).hexdigest()[:8]
    return f"{qhash}_{name}"

def infer_category(url, title, text=""):
    combined = norm(url + " " + title + " " + text[:1500])

    for category, keys in CATEGORY_RULES:
        if any(norm(k) in combined for k in keys):
            return category

    return "ostalo"

def get_sitemap_items():
    html = session.get(SITEMAP_URL, timeout=TIMEOUT).text
    soup = BeautifulSoup(html, "html.parser")

    items = []

    for a in soup.find_all("a", href=True):
        title = clean_text(a.get_text(" "))
        href = a.get("href", "")
        url = normalize_url(urljoin(SITEMAP_URL, href))

        if not url:
            continue

        # web.open pokazuje link IDs kroz redoslijed u dokumentu.
        # BeautifulSoup nema direktan ID, zato brojimo poziciju linka u listi.
        # Prvi link je index 0; u web prikazu relevantni su 4-180.
        # Zbog jezika na vrhu, ovo pouzdano odgovara sitemap strukturi.
        link_index = len(items)  # privremeno ne valja za filtriranje zbog skipova

    all_links = []

    for idx, a in enumerate(soup.find_all("a", href=True)):
        title = clean_text(a.get_text(" "))
        url = normalize_url(urljoin(SITEMAP_URL, a["href"]))

        if not url:
            continue

        if idx < MIN_LINK_ID or idx > MAX_LINK_ID:
            continue

        if should_drop(url, title):
            continue

        all_links.append({
            "url": url,
            "title": title,
            "sitemap_index": idx
        })

    seen = set()
    unique = []

    for item in all_links:
        if item["url"] not in seen:
            seen.add(item["url"])
            unique.append(item)

    return unique

def extract_html_text(html):
    try:
        extracted = trafilatura.extract(html)
        if extracted and len(extracted) > 80:
            return clean_text(extracted)
    except Exception:
        pass

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    return clean_text(soup.get_text(" "))

def extract_title(html, fallback):
    soup = BeautifulSoup(html, "html.parser")

    h1 = soup.find("h1")
    if h1:
        return clean_text(h1.get_text(" "))

    title = soup.find("title")
    if title:
        return clean_text(title.get_text(" "))

    return fallback

def extract_document_links(html, base_url):
    soup = BeautifulSoup(html, "html.parser")
    docs = []

    for a in soup.find_all("a", href=True):
        title = clean_text(a.get_text(" "))
        url = normalize_url(urljoin(base_url, a["href"]))

        if not url:
            continue

        if is_document_url(url):
            docs.append({"url": url, "title": title})

    return docs

def extract_pdf_text(path, max_pages=80):
    try:
        reader = PdfReader(path)
        parts = []

        for i, page in enumerate(reader.pages[:max_pages]):
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(f"[PAGE {i+1}] {txt}")

        return clean_text("\n".join(parts))
    except Exception:
        return ""

def extract_docx_text(path):
    try:
        doc = Document(path)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [clean_text(c.text) for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return clean_text("\n".join(parts))
    except Exception:
        return ""

def extract_excel_text(path, max_rows=300):
    try:
        excel = pd.ExcelFile(path)
        parts = []

        for sheet in excel.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet, nrows=max_rows)
            df = df.dropna(how="all")

            if not df.empty:
                parts.append(f"Sheet: {sheet}")
                parts.append(df.astype(str).to_csv(index=False))

        return clean_text("\n".join(parts))
    except Exception:
        return ""

def extract_file_text(path):
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf_text(path)

    if ext == ".docx":
        return extract_docx_text(path)

    if ext in [".xls", ".xlsx"]:
        return extract_excel_text(path)

    if ext in [".txt", ".rtf"]:
        try:
            return clean_text(path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            return ""

    return ""

def scrape_item(item):
    url = item["url"]
    fallback_title = item.get("title", "")

    try:
        r = session.get(url, timeout=TIMEOUT)

        if r.status_code != 200:
            return None, [], {"url": url, "status": "failed", "http_status": r.status_code}

        content_type = r.headers.get("Content-Type", "").lower()

        if is_document_url(url) or "pdf" in content_type or "officedocument" in content_type or "msword" in content_type or "spreadsheet" in content_type:
            filename = safe_filename(url)
            path = FILES_DIR / filename
            path.write_bytes(r.content)

            document_type = get_ext(url).replace(".", "") or "document"
            text = extract_file_text(path)
            title = fallback_title or filename
            category = infer_category(url, title, text)

            ocr_required = document_type == "pdf" and len(text) < 100

            record = {
                "source_url": url,
                "title": title,
                "category": category,
                "document_type": document_type,
                "content_type": content_type,
                "file_path": str(path),
                "text": text,
                "text_length": len(text),
                "ocr_required": ocr_required,
                "keep": True,
                "reason": "sitemap_document",
                "hash": hash_text(text) if text else ""
            }

            return record, [], {"url": url, "status": "ok", "http_status": 200}

        html = r.text
        title = extract_title(html, fallback_title)
        text = extract_html_text(html)
        docs = extract_document_links(html, url)
        category = infer_category(url, title, text)

        record = {
            "source_url": url,
            "title": title,
            "category": category,
            "document_type": "html",
            "content_type": content_type,
            "file_path": "",
            "text": text,
            "text_length": len(text),
            "linked_documents_count": len(docs),
            "ocr_required": False,
            "keep": len(text) >= MIN_TEXT_LENGTH or len(docs) > 0,
            "reason": "sitemap_page",
            "hash": hash_text(text) if text else ""
        }

        return record, docs, {"url": url, "status": "ok", "http_status": 200}

    except Exception as e:
        return None, [], {"url": url, "status": "error", "error": str(e)}

def split_sentences(text):
    return re.split(r"(?<=[.!?])\s+", clean_text(text))

def make_chunks(records):
    chunks = []
    seen = set()

    for r in records:
        if not r.get("keep") or r.get("ocr_required"):
            continue

        text = clean_text(r.get("text", ""))

        if len(text) < MIN_TEXT_LENGTH:
            continue

        sentences = split_sentences(text)
        current = ""
        idx = 0

        for sent in sentences:
            if len(current) + len(sent) <= CHUNK_SIZE:
                current += " " + sent
            else:
                chunk = clean_text(current)

                if len(chunk) >= MIN_TEXT_LENGTH:
                    h = hash_text(chunk)

                    if h not in seen:
                        seen.add(h)
                        chunks.append({
                            "chunk_id": f"{hash_text(r['source_url'])}_{idx}",
                            "source_url": r["source_url"],
                            "title": r["title"],
                            "category": r["category"],
                            "document_type": r["document_type"],
                            "chunk_index": idx,
                            "text": chunk
                        })
                        idx += 1

                current = chunk[-CHUNK_OVERLAP:] + " " + sent

        chunk = clean_text(current)

        if len(chunk) >= MIN_TEXT_LENGTH:
            h = hash_text(chunk)

            if h not in seen:
                seen.add(h)
                chunks.append({
                    "chunk_id": f"{hash_text(r['source_url'])}_{idx}",
                    "source_url": r["source_url"],
                    "title": r["title"],
                    "category": r["category"],
                    "document_type": r["document_type"],
                    "chunk_index": idx,
                    "text": chunk
                })

    return chunks

def save_jsonl(path, rows):
    with open(path, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")

def dedupe(records):
    seen_urls = set()
    seen_hashes = set()
    final = []

    for r in records:
        url = r["source_url"]
        h = r.get("hash", "")

        if url in seen_urls:
            continue

        if h and len(r.get("text", "")) > 250 and h in seen_hashes:
            continue

        seen_urls.add(url)

        if h:
            seen_hashes.add(h)

        final.append(r)

    return final

def main():
    sitemap_items = get_sitemap_items()
    print("Sitemap items selected:", len(sitemap_items))

    pd.DataFrame(sitemap_items).to_csv(OUTPUT_DIR / "selected_sitemap_links.csv", index=False)

    records = []
    events = []
    doc_items = []

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = [ex.submit(scrape_item, item) for item in sitemap_items]

        for fut in tqdm(as_completed(futures), total=len(futures), desc="Scraping sitemap pages"):
            record, docs, event = fut.result()
            events.append(event)

            if record:
                records.append(record)

            doc_items.extend(docs)

    seen_docs = set()
    unique_docs = []

    for d in doc_items:
        if d["url"] not in seen_docs and not should_drop(d["url"], d.get("title", "")):
            seen_docs.add(d["url"])
            unique_docs.append(d)

    print("Document links found:", len(unique_docs))

    if unique_docs:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
            futures = [ex.submit(scrape_item, item) for item in unique_docs]

            for fut in tqdm(as_completed(futures), total=len(futures), desc="Scraping documents"):
                record, _, event = fut.result()
                events.append(event)

                if record:
                    records.append(record)

    records = dedupe(records)
    chunks = make_chunks(records)

    df = pd.DataFrame(records)
    chunks_df = pd.DataFrame(chunks)
    events_df = pd.DataFrame(events)

    df.to_csv(OUTPUT_DIR / "all_records_v3_fixed.csv", index=False)
    df[df["keep"] == True].to_csv(OUTPUT_DIR / "kept_records_v3_fixed.csv", index=False)
    df[df["ocr_required"] == True].to_csv(OUTPUT_DIR / "needs_ocr_v3_fixed.csv", index=False)
    chunks_df.to_csv(OUTPUT_DIR / "chunks_v3_fixed.csv", index=False)
    events_df.to_csv(OUTPUT_DIR / "crawl_events_v3_fixed.csv", index=False)

    save_jsonl(OUTPUT_DIR / "chunks_v3_fixed.jsonl", chunks)

    coverage = {
        "selected_sitemap_links": len(sitemap_items),
        "document_links_found": len(unique_docs),
        "total_records": len(df),
        "kept_records": int((df["keep"] == True).sum()) if not df.empty else 0,
        "total_chunks": len(chunks),
        "ocr_required": int((df["ocr_required"] == True).sum()) if not df.empty else 0,
        "records_by_category": df[df["keep"] == True]["category"].value_counts().to_dict() if not df.empty else {},
        "chunks_by_category": chunks_df["category"].value_counts().to_dict() if not chunks_df.empty else {},
        "records_by_type": df[df["keep"] == True]["document_type"].value_counts().to_dict() if not df.empty else {},
    }

    with open(OUTPUT_DIR / "coverage_report_v3_fixed.json", "w", encoding="utf-8") as f:
        json.dump(coverage, f, indent=2, ensure_ascii=False)

    print("\nDONE")
    print(json.dumps(coverage, indent=2, ensure_ascii=False))
    print("\nCheck selected links:", OUTPUT_DIR / "selected_sitemap_links.csv")

if __name__ == "__main__":
    main()