import os
import re
import json
import time
import hashlib
import logging
from pathlib import Path
from urllib.parse import urljoin, urlparse, urldefrag
from collections import deque, Counter
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
import pandas as pd
import matplotlib.pyplot as plt
import trafilatura

from bs4 import BeautifulSoup
from tqdm import tqdm
from pypdf import PdfReader
from docx import Document
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


# ============================
# CONFIG
# ============================

OUTPUT_DIR = Path("mpr_dataset_v2")
FILES_DIR = OUTPUT_DIR / "files"
REPORTS_DIR = OUTPUT_DIR / "reports"

FILES_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_DOMAIN = "mpr.gov.ba"

MAX_PAGES = 2500
MAX_WORKERS = 8
TIMEOUT = 25
CHECKPOINT_EVERY = 100

MIN_TEXT_LENGTH = 180
CHUNK_SIZE = 1100
CHUNK_OVERLAP = 180

HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; MPR Dataset Scraper v2; student-project)"
}

DOCUMENT_EXTENSIONS = {
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".rtf"
}

START_URLS = [
    # nova struktura
    "https://mpr.gov.ba/bs/mapa-stranica",
    "https://mpr.gov.ba/bs/nvo",
    "https://mpr.gov.ba/bs/pravosudni-ispit",
    "https://mpr.gov.ba/bs/strucni-upravni-ispit-za-srednju-i-visu-skolsku-spremu",
    "https://mpr.gov.ba/bs/besplatna-pravna-pomoc",
    "https://mpr.gov.ba/bs/potrebna-dokumentacija",
    "https://mpr.gov.ba/bs/kako-osnovati-udruzenje",
    "https://mpr.gov.ba/bs/registracija-fondacije",
    "https://mpr.gov.ba/bs/administrativne-takse87",

    # stara struktura - jako bitno za obrasce
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/udruzenja/obrasci/?id=9974",
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/fondacije/obrasci/?id=2104",
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/upis_promjena/obrasci/?id=8536",
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/udruzenja/",
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/fondacije/",
    "https://www.mpr.gov.ba/organizacija_nadleznosti/uprava/registracije/",
]

DROP_URL_PATTERNS = [
    r"/hr($|/)",
    r"/en($|/)",
    r"/sr($|/)",
    r"/portal/objava/",
    r"/portal\?",
    r"categorySlug=vijesti",
    r"vijesti",
    r"novosti",
    r"saopstenja",
    r"saopštenja",
    r"tender",
    r"javne-nabavke",
    r"javni-oglasi",
    r"konkurs",
    r"galerija",
    r"foto",
    r"video",
    r"eu4justice",
    r"projekti",
    r"strategije",
    r"transparentnost",
    r"oglasna-ploca",
    r"oglasna-ploča",
]

CATEGORY_RULES = {
    "obrasci": [
        "obrasci", "obrazac", "formular", "zahtjev", "prijava",
        "obrazac-1", "obrazac_1"
    ],
    "registracija": [
        "registracija", "registracije", "udruzenja", "udruženja",
        "fondacije", "nvo", "kako-osnovati", "potrebna-dokumentacija",
        "administrativne-takse", "upis-promjena", "brisanja-iz-registra",
        "registar-udruzenja", "registar-udruženja"
    ],
    "ispiti": [
        "pravosudni-ispit", "pravosudni ispit",
        "strucni-upravni-ispit", "stručni-upravni-ispit",
        "strucni upravni ispit", "stručni upravni ispit",
        "ispitni-termini", "novi-ispitni-termini"
    ],
    "pravna_pomoc": [
        "besplatna-pravna-pomoc", "besplatna-pravna-pomoć",
        "pravna-pomoc", "pravna-pomoć",
        "medjunarodna-pravna-pomoc", "međunarodna-pravna-pomoć",
        "alimentacija", "otmica-djece", "otmice-djece"
    ],
    "zakoni_i_propisi": [
        "zakoni", "zakon", "propisi", "pravilnici",
        "podzakonski", "biblioteka-zakona", "sluzbeni-glasnik",
        "službeni-glasnik", "konvencije", "ugovori"
    ],
    "registri": [
        "registar", "registri", "evidencija", "lista"
    ],
    "notari": [
        "notari", "notar", "notarsk"
    ],
    "sudski_tumaci": [
        "sudski-tumaci", "sudski-tumači",
        "sudski tumaci", "sudski tumači",
        "stalni-sudski", "tumac", "tumač", "tumaci", "tumači"
    ],
    "kontakt_i_nadleznosti": [
        "kontakt", "kontakti", "organizacija", "nadleznosti",
        "nadležnosti", "sektor", "ministarstvo"
    ],
}

CATEGORY_PRIORITY = [
    "sudski_tumaci",
    "notari",
    "obrasci",
    "registracija",
    "ispiti",
    "pravna_pomoc",
    "zakoni_i_propisi",
    "registri",
    "kontakt_i_nadleznosti",
]

RELEVANT_TERMS = [
    "zahtjev", "obrazac", "registracija", "preregistracija",
    "udruženje", "udruzenje", "fondacija", "zakon", "pravilnik",
    "propis", "pravna pomoć", "pravna pomoc", "notar", "sudski tumač",
    "sudski tumac", "pravosudni ispit", "stručni upravni ispit",
    "strucni upravni ispit", "registar", "taksa", "dokumentacija",
    "nadležnost", "nadleznost", "rok", "rješenje", "rjesenje",
    "ministarstvo pravde", "formular", "prijava"
]

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")


# ============================
# SESSION
# ============================

def make_session():
    session = requests.Session()
    session.headers.update(HEADERS)

    retries = Retry(
        total=3,
        backoff_factor=0.6,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"]
    )

    adapter = HTTPAdapter(max_retries=retries, pool_connections=50, pool_maxsize=50)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    return session


# ============================
# HELPERS
# ============================

def clean_text(text):
    if not text:
        return ""
    text = str(text).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_text(text):
    text = clean_text(text).lower()
    text = text.replace("č", "c").replace("ć", "c")
    text = text.replace("š", "s").replace("ž", "z").replace("đ", "dj")
    return text


def normalize_url(url):
    if not url:
        return None

    url, _ = urldefrag(url)
    parsed = urlparse(url)

    if parsed.scheme not in ("http", "https"):
        return None

    if not parsed.netloc.endswith(ALLOWED_DOMAIN):
        return None

    return url.rstrip("/")


def should_drop_url(url):
    n = normalize_text(url)
    return any(re.search(pattern, n, flags=re.IGNORECASE) for pattern in DROP_URL_PATTERNS)


def is_document_url(url):
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in DOCUMENT_EXTENSIONS)


def get_extension(url):
    return Path(urlparse(url).path).suffix.lower()


def safe_filename(url):
    parsed = urlparse(url)
    name = os.path.basename(parsed.path)

    if not name:
        name = hashlib.md5(url.encode()).hexdigest() + ".html"

    name = re.sub(r"[^a-zA-Z0-9_.-]", "_", name)

    qhash = hashlib.md5((parsed.query or "").encode()).hexdigest()[:8]

    if "." not in name:
        name += ".html"

    return f"{qhash}_{name}"


def hash_text(text):
    return hashlib.md5(clean_text(text).encode("utf-8")).hexdigest()


def infer_category(url, title="", text="", parent_title="", anchor_text=""):
    combined = normalize_text(" ".join([
        url or "",
        title or "",
        parent_title or "",
        anchor_text or "",
        text[:1500] if text else ""
    ]))

    for category in CATEGORY_PRIORITY:
        terms = CATEGORY_RULES[category]
        if any(normalize_text(term) in combined for term in terms):
            return category

    return "ostalo"


def has_relevant_terms(text):
    n = normalize_text(text)
    return any(normalize_text(term) in n for term in RELEVANT_TERMS)


def title_from_soup(soup, fallback=""):
    h1 = soup.find("h1")
    if h1:
        return clean_text(h1.get_text(" "))

    title = soup.find("title")
    if title:
        return clean_text(title.get_text(" "))

    return fallback


def extract_html_text(html):
    try:
        extracted = trafilatura.extract(html)
        if extracted and len(extracted) > 100:
            return clean_text(extracted)
    except Exception:
        pass

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    return clean_text(soup.get_text(" "))


def extract_links(html, base_url):
    soup = BeautifulSoup(html, "html.parser")
    results = []

    for a in soup.find_all("a", href=True):
        raw = a["href"]
        text = clean_text(a.get_text(" "))
        full = urljoin(base_url, raw)
        full = normalize_url(full)

        if not full:
            continue

        results.append({
            "url": full,
            "anchor_text": text
        })

    return results


# ============================
# FILE EXTRACTION
# ============================

def extract_pdf_text(path, max_pages=50):
    try:
        reader = PdfReader(path)
        pages = []

        for i, page in enumerate(reader.pages[:max_pages]):
            txt = page.extract_text() or ""
            if txt.strip():
                pages.append(f"[PAGE {i + 1}] {txt}")

        return clean_text("\n".join(pages))
    except Exception:
        return ""


def extract_docx_text(path):
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return clean_text("\n".join(parts))
    except Exception:
        return ""


def extract_excel_text(path, max_rows=200):
    try:
        excel = pd.ExcelFile(path)
        parts = []

        for sheet in excel.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet, nrows=max_rows)
            df = df.dropna(how="all")

            if df.empty:
                continue

            parts.append(f"Sheet: {sheet}")
            parts.append(df.astype(str).to_csv(index=False))

        return clean_text("\n".join(parts))
    except Exception:
        return ""


def extract_file_text(path):
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf_text(path)

    if ext == ".docx":
        return extract_docx_text(path)

    if ext in [".xls", ".xlsx"]:
        return extract_excel_text(path)

    if ext in [".txt", ".rtf"]:
        try:
            return clean_text(path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            return ""

    return ""


# ============================
# CLASSIFICATION
# ============================

def decide_keep(url, title, text, category, document_type, linked_docs_count=0):
    if should_drop_url(url):
        return False, "dropped_by_url_filter", "irrelevant"

    if document_type in ["pdf", "doc", "docx", "xls", "xlsx"]:
        if category != "ostalo":
            return True, "kept_relevant_document", "ok"

        if has_relevant_terms(title + " " + text):
            return True, "kept_document_relevant_terms", "ok"

        return False, "dropped_document_not_relevant", "needs_review"

    if linked_docs_count > 0 and category != "ostalo":
        return True, "kept_index_page_with_documents", "index_page"

    if len(clean_text(text)) < MIN_TEXT_LENGTH:
        return False, "dropped_too_short", "too_short"

    if category != "ostalo":
        return True, "kept_relevant_category", "ok"

    if has_relevant_terms(title + " " + text):
        return True, "kept_relevant_terms", "needs_review"

    return False, "dropped_not_relevant", "irrelevant"


# ============================
# SCRAPE TASK
# ============================

def scrape_one(task):
    session = make_session()

    url = task["url"]
    parent_url = task.get("parent_url", "")
    parent_title = task.get("parent_title", "")
    anchor_text = task.get("anchor_text", "")

    try:
        response = session.get(url, timeout=TIMEOUT)

        event = {
            "url": url,
            "status": "ok",
            "http_status": response.status_code,
            "error": ""
        }

        if response.status_code != 200:
            event["status"] = "failed"
            return None, [], event

        content_type = response.headers.get("Content-Type", "").lower()
        ext = get_extension(url)

        is_doc = (
            is_document_url(url)
            or "application/pdf" in content_type
            or "msword" in content_type
            or "officedocument" in content_type
            or "spreadsheet" in content_type
        )

        if is_doc:
            filename = safe_filename(url)
            file_path = FILES_DIR / filename
            file_path.write_bytes(response.content)

            document_type = ext.replace(".", "") if ext else "document"
            text = extract_file_text(file_path)
            title = anchor_text or filename

            ocr_required = document_type == "pdf" and len(text) < 100
            category = infer_category(url, title, text, parent_title, anchor_text)

            keep, reason, quality_flag = decide_keep(
                url=url,
                title=title,
                text=text,
                category=category,
                document_type=document_type,
                linked_docs_count=0
            )

            if ocr_required and category != "ostalo":
                keep = True
                reason = "kept_relevant_pdf_needs_ocr"
                quality_flag = "needs_ocr"

            record = {
                "source_url": url,
                "title": title,
                "category": category,
                "document_type": document_type,
                "content_type": content_type,
                "file_path": str(file_path),
                "parent_url": parent_url,
                "parent_title": parent_title,
                "anchor_text": anchor_text,
                "text": text,
                "text_length": len(text),
                "linked_documents_count": 0,
                "ocr_required": ocr_required,
                "processing_status": "needs_ocr" if ocr_required else "text_extracted",
                "keep": keep,
                "reason": reason,
                "quality_flag": quality_flag,
                "hash": hash_text(text) if text else ""
            }

            return record, [], event

        html = response.text
        soup = BeautifulSoup(html, "html.parser")

        title = title_from_soup(soup, fallback=anchor_text or url)
        text = extract_html_text(html)
        links = extract_links(html, url)

        linked_docs_count = sum(1 for x in links if is_document_url(x["url"]))
        category = infer_category(url, title, text, parent_title, anchor_text)

        keep, reason, quality_flag = decide_keep(
            url=url,
            title=title,
            text=text,
            category=category,
            document_type="html",
            linked_docs_count=linked_docs_count
        )

        record = {
            "source_url": url,
            "title": title,
            "category": category,
            "document_type": "html",
            "content_type": content_type,
            "file_path": "",
            "parent_url": parent_url,
            "parent_title": parent_title,
            "anchor_text": anchor_text,
            "text": text,
            "text_length": len(text),
            "linked_documents_count": linked_docs_count,
            "ocr_required": False,
            "processing_status": "text_extracted",
            "keep": keep,
            "reason": reason,
            "quality_flag": quality_flag,
            "hash": hash_text(text) if text else ""
        }

        next_tasks = []

        for link in links:
            link_url = link["url"]

            if should_drop_url(link_url):
                continue

            # Prati dokumente uvijek ako link dolazi iz relevantne parent stranice
            # Prati HTML samo ako izgleda relevantno
            if is_document_url(link_url):
                next_tasks.append({
                    "url": link_url,
                    "parent_url": url,
                    "parent_title": title,
                    "anchor_text": link["anchor_text"]
                })
            else:
                link_category = infer_category(
                    link_url,
                    link.get("anchor_text", ""),
                    "",
                    title,
                    link.get("anchor_text", "")
                )

                if link_category != "ostalo" or category != "ostalo":
                    next_tasks.append({
                        "url": link_url,
                        "parent_url": url,
                        "parent_title": title,
                        "anchor_text": link["anchor_text"]
                    })

        return record, next_tasks, event

    except Exception as e:
        return None, [], {
            "url": url,
            "status": "error",
            "http_status": "",
            "error": str(e)
        }


# ============================
# CHUNKING
# ============================

def split_into_sentences(text):
    text = clean_text(text)
    return re.split(r"(?<=[.!?])\s+", text)


def make_chunks(records):
    chunks = []
    seen = set()

    for r in records:
        if not r.get("keep"):
            continue

        if r.get("ocr_required"):
            continue

        text = clean_text(r.get("text", ""))

        if len(text) < MIN_TEXT_LENGTH:
            continue

        sentences = split_into_sentences(text)

        current = ""
        chunk_index = 0

        for sent in sentences:
            if len(current) + len(sent) <= CHUNK_SIZE:
                current += " " + sent
            else:
                chunk = clean_text(current)

                if len(chunk) >= MIN_TEXT_LENGTH:
                    h = hash_text(chunk)
                    if h not in seen:
                        seen.add(h)
                        chunks.append({
                            "chunk_id": f"{hash_text(r['source_url'])}_{chunk_index}",
                            "source_url": r["source_url"],
                            "title": r["title"],
                            "category": r["category"],
                            "document_type": r["document_type"],
                            "parent_url": r.get("parent_url", ""),
                            "parent_title": r.get("parent_title", ""),
                            "chunk_index": chunk_index,
                            "text": chunk
                        })
                        chunk_index += 1

                # overlap: zadrži kraj prethodnog chunka
                current = chunk[-CHUNK_OVERLAP:] + " " + sent

        final_chunk = clean_text(current)

        if len(final_chunk) >= MIN_TEXT_LENGTH:
            h = hash_text(final_chunk)
            if h not in seen:
                seen.add(h)
                chunks.append({
                    "chunk_id": f"{hash_text(r['source_url'])}_{chunk_index}",
                    "source_url": r["source_url"],
                    "title": r["title"],
                    "category": r["category"],
                    "document_type": r["document_type"],
                    "parent_url": r.get("parent_url", ""),
                    "parent_title": r.get("parent_title", ""),
                    "chunk_index": chunk_index,
                    "text": final_chunk
                })

    return chunks


# ============================
# OUTPUTS
# ============================

def save_jsonl(path, rows):
    with open(path, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def deduplicate_records(records):
    result = []
    seen_url = set()
    seen_hash = set()

    for r in records:
        url = r.get("source_url", "")
        h = r.get("hash", "")

        if url in seen_url:
            r["keep"] = False
            r["reason"] = "duplicate_url"
            r["quality_flag"] = "duplicate"

        elif h and h in seen_hash and len(r.get("text", "")) > 100:
            r["keep"] = False
            r["reason"] = "duplicate_content"
            r["quality_flag"] = "duplicate"

        seen_url.add(url)

        if h:
            seen_hash.add(h)

        result.append(r)

    return result


def generate_reports(records, chunks, events):
    df = pd.DataFrame(records)
    chunks_df = pd.DataFrame(chunks)
    events_df = pd.DataFrame(events)

    df.to_csv(OUTPUT_DIR / "all_records_v2.csv", index=False)
    df[df["keep"] == True].to_csv(OUTPUT_DIR / "kept_records_v2.csv", index=False)
    df[df["keep"] == False].to_csv(OUTPUT_DIR / "dropped_records_v2.csv", index=False)
    df[df["ocr_required"] == True].to_csv(OUTPUT_DIR / "needs_ocr_v2.csv", index=False)
    df[df["quality_flag"].isin(["needs_review", "index_page"])].to_csv(OUTPUT_DIR / "needs_review_v2.csv", index=False)

    save_jsonl(OUTPUT_DIR / "chunks_v2.jsonl", chunks)
    chunks_df.to_csv(OUTPUT_DIR / "chunks_v2.csv", index=False)
    events_df.to_csv(OUTPUT_DIR / "crawl_events_v2.csv", index=False)

    coverage = {
        "records_by_category": df[df["keep"] == True]["category"].value_counts().to_dict(),
        "chunks_by_category": chunks_df["category"].value_counts().to_dict() if not chunks_df.empty else {},
        "records_by_type": df[df["keep"] == True]["document_type"].value_counts().to_dict(),
        "ocr_required": int((df["ocr_required"] == True).sum()),
        "total_records": int(len(df)),
        "kept_records": int((df["keep"] == True).sum()),
        "dropped_records": int((df["keep"] == False).sum()),
        "total_chunks": int(len(chunks)),
    }

    with open(OUTPUT_DIR / "coverage_report_v2.json", "w", encoding="utf-8") as f:
        json.dump(coverage, f, indent=2, ensure_ascii=False)

    def bar(series, title, filename):
        if series.empty:
            return
        plt.figure(figsize=(11, 5))
        series.plot(kind="bar")
        plt.title(title)
        plt.ylabel("Broj")
        plt.tight_layout()
        plt.savefig(REPORTS_DIR / filename, dpi=150)
        plt.close()

    if not df.empty:
        bar(df["keep"].value_counts(), "Keep vs Drop", "keep_drop_v2.png")
        bar(df[df["keep"] == True]["category"].value_counts(), "Kept records by category", "kept_categories_v2.png")
        bar(df["document_type"].value_counts(), "Document types", "document_types_v2.png")
        bar(df["reason"].value_counts().head(20), "Decision reasons", "reasons_v2.png")

    if not chunks_df.empty:
        bar(chunks_df["category"].value_counts(), "Chunks by category", "chunks_categories_v2.png")

    html = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <title>MPR Dataset V2 Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; background: #fafafa; }}
            .card {{ background: white; padding: 20px; margin-bottom: 20px; border-radius: 12px; box-shadow: 0 2px 10px rgba(0,0,0,.08); }}
            img {{ max-width: 950px; width: 100%; }}
            code {{ background: #eee; padding: 2px 6px; border-radius: 4px; }}
            table {{ border-collapse: collapse; width: 100%; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; }}
            th {{ background: #eee; }}
        </style>
    </head>
    <body>
        <h1>MPR Dataset V2 Report</h1>

        <div class="card">
            <h2>Summary</h2>
            <table>
                <tr><th>Total records</th><td>{coverage["total_records"]}</td></tr>
                <tr><th>Kept records</th><td>{coverage["kept_records"]}</td></tr>
                <tr><th>Dropped records</th><td>{coverage["dropped_records"]}</td></tr>
                <tr><th>Total chunks</th><td>{coverage["total_chunks"]}</td></tr>
                <tr><th>OCR required</th><td>{coverage["ocr_required"]}</td></tr>
            </table>
        </div>

        <div class="card">
            <h2>Files</h2>
            <ul>
                <li><code>all_records_v2.csv</code></li>
                <li><code>kept_records_v2.csv</code></li>
                <li><code>dropped_records_v2.csv</code></li>
                <li><code>needs_review_v2.csv</code></li>
                <li><code>needs_ocr_v2.csv</code></li>
                <li><code>chunks_v2.jsonl</code></li>
                <li><code>chunks_v2.csv</code></li>
                <li><code>coverage_report_v2.json</code></li>
            </ul>
        </div>

        {''.join(f'<div class="card"><img src="reports/{p.name}"></div>' for p in REPORTS_DIR.glob("*_v2.png"))}
    </body>
    </html>
    """

    with open(OUTPUT_DIR / "report_v2.html", "w", encoding="utf-8") as f:
        f.write(html)

    print("\nDONE")
    print(json.dumps(coverage, indent=2, ensure_ascii=False))
    print(f"\nOpen report: {OUTPUT_DIR / 'report_v2.html'}")


def checkpoint(records, events):
    pd.DataFrame(records).to_csv(OUTPUT_DIR / "checkpoint_records_v2.csv", index=False)
    pd.DataFrame(events).to_csv(OUTPUT_DIR / "checkpoint_events_v2.csv", index=False)


# ============================
# CRAWLER
# ============================

def crawl():
    queue = deque()
    queued = set()
    visited = set()
    records = []
    events = []

    for url in START_URLS:
        u = normalize_url(url)
        if u and not should_drop_url(u):
            queue.append({
                "url": u,
                "parent_url": "",
                "parent_title": "",
                "anchor_text": ""
            })
            queued.add(u)

    pbar = tqdm(total=MAX_PAGES, desc="MPR scraper v2", unit="url")

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        while queue and len(visited) < MAX_PAGES:
            batch = []

            while queue and len(batch) < MAX_WORKERS * 3 and len(visited) + len(batch) < MAX_PAGES:
                task = queue.popleft()

                if task["url"] in visited:
                    continue

                visited.add(task["url"])
                batch.append(task)

            futures = {executor.submit(scrape_one, task): task for task in batch}

            for future in as_completed(futures):
                record, next_tasks, event = future.result()

                events.append(event)

                if record:
                    records.append(record)

                for nt in next_tasks:
                    u = nt["url"]

                    if u not in visited and u not in queued and not should_drop_url(u):
                        queued.add(u)
                        queue.append(nt)

                pbar.update(1)

                kept = sum(1 for r in records if r.get("keep"))
                docs = sum(1 for r in records if r.get("document_type") != "html")
                ocr = sum(1 for r in records if r.get("ocr_required"))

                pbar.set_postfix({
                    "records": len(records),
                    "kept": kept,
                    "docs": docs,
                    "ocr": ocr,
                    "queue": len(queue)
                })

                if len(records) > 0 and len(records) % CHECKPOINT_EVERY == 0:
                    checkpoint(records, events)

    pbar.close()

    records = deduplicate_records(records)
    chunks = make_chunks(records)
    generate_reports(records, chunks, events)


if __name__ == "__main__":
    crawl()